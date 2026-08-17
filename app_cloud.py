import streamlit as st
import streamlit.components.v1 as components
import os
import sys
import json
import tempfile
import requests
import base64
import gc
import copy
import string
import shutil
import io
import re
from datetime import datetime, timedelta
import soundfile as sf
import numpy as np
import google.generativeai as genai
from my_custom_editor.my_component import custom_transcription_editor
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ==========================================
# פונקציות תמלול ו-AI
# ==========================================
def diarize_with_deepgram(api_key, file_path, expected_speakers=0):
    """קריאה ל-Deepgram לזיהוי דוברים (מהיר וחינמי בהרשמה)"""
    url = "https://api.deepgram.com/v1/listen?diarize=true"
    if expected_speakers > 0:
        pass # Deepgram מטפל בזה אוטומטית בצורה טובה, אפשר להשאיר ריק
        
    headers = {
        "Authorization": f"Token {api_key}",
        "Content-Type": "audio/wav"
    }
    with open(file_path, "rb") as f:
        response = requests.post(url, headers=headers, data=f)
        
    if response.status_code != 200:
        raise Exception(f"שגיאת API: {response.text}")
        
    data = response.json()
    words = data['results']['channels'][0]['alternatives'][0]['words']
    
    # המרה לפורמט הסטנדרטי של האפליקציה שלך
    turns = []
    for w in words:
        if 'speaker' in w:
            turns.append({"start": w['start'], "end": w['end'], "speaker": str(w['speaker'])})
    return turns

def diarize_with_replicate(api_key, file_path, expected_speakers=0):
    """קריאה ל-Pyannote דרך Replicate API (דיוק מקסימלי, ענן)"""
    url = "https://api.replicate.com/v1/predictions"
    headers = {"Authorization": f"Token {api_key}", "Content-Type": "application/json"}
    
    # קריאת הקובץ והמרה ל-Base64 כדי לשלוח ל-API
    with open(file_path, "rb") as f:
        audio_base64 = base64.b64encode(f.read()).decode("utf-8")
        audio_uri = f"data:audio/wav;base64,{audio_base64}"

    data = {
        "version": "3a0778f6570c0c08283d03823758b90c1e550c61bba104a3206411ddfc278a3f", # Pyannote 3.1
        "input": {"audio": audio_uri}
    }
    if expected_speakers > 0:
        data["input"]["num_speakers"] = expected_speakers

    response = requests.post(url, headers=headers, json=data)
    if response.status_code != 201:
        raise Exception(f"שגיאת API: {response.text}")
        
    prediction_url = response.json()["urls"]["get"]
    
    # המתנה לתוצאה מהענן
    import time
    while True:
        poll = requests.get(prediction_url, headers=headers).json()
        if poll["status"] == "succeeded":
            return poll["output"]["segments"]
        elif poll["status"] == "failed":
            raise Exception("הזיהוי נכשל בשרת")
        time.sleep(2)

def transcribe_with_groq_api(api_key, file_path, language, initial_prompt=""):
    """שליחת האודיו ל-API של Groq והמרת התוצאה לפורמט המילים של המערכת"""
    url = "https://api.groq.com/openai/v1/audio/transcriptions"
    headers = {"Authorization": f"Bearer {api_key}"}
    
    with open(file_path, "rb") as f:
        files = {"file": (os.path.basename(file_path), f, "audio/wav")}
        data = {
            "model": "whisper-large-v3",
            "response_format": "verbose_json",
            "language": language,
            "timestamp_granularities[]": "segment" 
        }
        if initial_prompt:
            data["prompt"] = initial_prompt
            
        response = requests.post(url, headers=headers, files=files, data=data)
        
        if response.status_code != 200:
            raise Exception(f"שגיאת API: {response.text}")
            
        api_json = response.json()
        
    words_list = []
    word_id = 0
    
    # חלוקה גסה של זמן המקטעים למילים 
    for seg in api_json.get("segments", []):
        seg_text = seg.get("text", "").strip().split()
        if not seg_text: continue
        
        seg_start = seg.get("start", 0.0)
        seg_end = seg.get("end", seg_start + 1.0)
        chunk_time = (seg_end - seg_start) / max(len(seg_text), 1)
        
        for i, raw_word in enumerate(seg_text):
            clean_w = raw_word.strip(string.punctuation + '.,?!״׳"()[]{}')
            prefix_punc = raw_word[:raw_word.find(clean_w)] if clean_w else ""
            suffix_punc = raw_word[raw_word.find(clean_w) + len(clean_w):] if clean_w else raw_word
            
            words_list.append({
                "id": word_id, 
                "word": clean_w, 
                "prefix_punc": prefix_punc, 
                "punctuation": suffix_punc,
                "original_word": raw_word, 
                "clean_word": clean_w.lower(), 
                "start": seg_start + i * chunk_time,
                "end": seg_start + (i + 1) * chunk_time, 
                "confidence": 0.99, 
                "speaker": "0", 
                "deleted": False
            })
            word_id += 1
            
    return words_list, api_json.get("duration", 0.0)

def transcribe_with_premium_api(api_key, file_path, language, initial_prompt=""):
    """
    מסלול פרימיום: הכנה לשלב 2. 
    יוחלף בהמשך ל-ivrit.ai או Deepgram שיספקו ציוני ביטחון לכל מילה.
    """
    st.toast("משתמש במנוע הזמני עד חיבור ה-API של הפרימיום (שלב 2)")
    return transcribe_with_groq_api(api_key, file_path, language, initial_prompt)

def call_gemini_api(api_key, text, prompt_type, custom_prompt=""):
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')

        if prompt_type == "סיכום כללי":
            system_prompt = "אתה עוזר וירטואלי מקצועי בעברית. סכם את התמלול הבא בצורה ברורה, קולחת ותמציתית."
        elif prompt_type == "נקודות מפתח (Bullet points)":
            system_prompt = "אתה עוזר וירטואלי מקצועי בעברית. חלץ את נקודות המפתח והנושאים המרכזיים מהתמלול הבא, והצג אותם כרשימה (Bullet points) מסודרת."
        elif prompt_type == "ניתוח (אווירה ומסקנות)":
            system_prompt = "אתה מנתח שיחות מקצועי. קרא את התמלול הבא וכתוב: 1. מהי האווירה הכללית והדינמיקה בשיחה? 2. מהן המסקנות העיקריות או הפעולות להמשך (Action items) שעולות ממנה?"
        else: 
            system_prompt = custom_prompt

        full_prompt = f"{system_prompt}\n\nהנה התמלול:\n{text}"
        response = model.generate_content(full_prompt)
        return response.text
    except Exception as e:
        return f"❌ שגיאה בתקשורת עם AI: {str(e)}\n\n(וודא שמפתח ה-API שלך תקין)"

def auto_punctuate_with_gemini(api_key, words_data):
    """פונקציה ששולחת את המילים ל-AI ומקבלת חזרה אובייקט JSON עם סימני פיסוק בלבד כדי לא לפגוע בתזמונים"""
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        active_words = [w for w in words_data if not w.get("deleted")]
        numbered_text = "\n".join([f"{w['id']}: {w['word']}" for w in active_words])
        
        prompt = f"""
        אתה עורך לשוני מקצועי. קיבלת רשימה של מילים מתוך תמלול, ממוספרות לפי מזהה (ID).
        המטרה שלך היא להחליט אילו מילים חייבות סימן פיסוק מיד אחריהן (פסיק, נקודה, או סימן שאלה) כדי ליצור חלוקה הגיונית למשפטים.
        
        החזר אך ורק אובייקט JSON חוקי! 
        המפתח יהיה ה-ID של המילה, והערך יהיה סימן הפיסוק (",", ".", "?").
        אל תחזיר אף מילה שאין לה סימן פיסוק!
        דוגמה לתשובה תקינה: {{"5": ",", "12": ".", "24": "?"}}
        
        רשימת המילים:
        {numbered_text}
        """
        
        response = model.generate_content(
            prompt,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json"
            )
        )
        return json.loads(response.text)
    except Exception as e:
        print(f"Error in Gemini JSON generation: {e}")
        return None

# ==========================================
# הגדרות אפליקציה ותיקיות
# ==========================================
st.set_page_config(page_title="Advanced STT - Cloud Editor", layout="wide")

PROJECTS_DIR = "saved_transcriptions"
TEMP_DIR = "temp_uploads"
os.makedirs(PROJECTS_DIR, exist_ok=True)
os.makedirs(TEMP_DIR, exist_ok=True)

st.title("🎙️ מערכת תמלול ועריכה - גרסת ענן")
st.markdown("תמלול מהיר בענן, ניהול דוברים אינטראקטיבי ותיקון שגיאות חכם.")

# ==========================================
# ניהול מצב (State)
# ==========================================
if "words_data" not in st.session_state: st.session_state.words_data = []
if "original_words_data" not in st.session_state: st.session_state.original_words_data = []
if "audio_bytes" not in st.session_state: st.session_state.audio_bytes = None
if "current_file_name" not in st.session_state: st.session_state.current_file_name = None
if "active_word_id" not in st.session_state: st.session_state.active_word_id = "" 
if "speaker_names" not in st.session_state: st.session_state.speaker_names = {}
if "project_dir" not in st.session_state: st.session_state.project_dir = None
if "local_audio_path" not in st.session_state: st.session_state.local_audio_path = None
if "decoded_audio" not in st.session_state: st.session_state.decoded_audio = None
if "history" not in st.session_state: st.session_state.history = []
if "history_index" not in st.session_state: st.session_state.history_index = -1

FILLERS = ["אה", "אממ", "אהה", "אמ", "um", "uh", "mhm", "mm", "ah", "er", "hmm"]

# ==========================================
# פונקציות עזר 
# ==========================================
def get_word_color(confidence):
    if confidence < 0.5: return "#ff4b4b" 
    elif confidence <= 0.9: return "#ffe14b" 
    else: return "#f0f2f6" 

def commit_to_history():
    if st.session_state.history_index < len(st.session_state.history) - 1:
        st.session_state.history = st.session_state.history[:st.session_state.history_index + 1]
    st.session_state.history.append(copy.deepcopy(st.session_state.words_data))
    st.session_state.history_index = len(st.session_state.history) - 1

def save_project():
    if st.session_state.project_dir:
        data = {
            "original_words_data": st.session_state.original_words_data,
            "words_data": st.session_state.words_data,
            "speaker_names": st.session_state.speaker_names
        }
        json_path = os.path.join(st.session_state.project_dir, "data.json")
        temp_json_path = os.path.join(st.session_state.project_dir, "data.json.tmp")
        try:
            with open(temp_json_path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            os.replace(temp_json_path, json_path)
        except Exception as e:
            st.toast(f"שגיאה בשמירת הפרויקט: {e}")

def slice_audio(start_sec, end_sec, padding=1.5):
    if not st.session_state.get("local_audio_path"): return None
    try:
        data, sr = sf.read(st.session_state.local_audio_path)
        if len(data.shape) > 1: data = np.mean(data, axis=1)
        
        start_time = max(0, start_sec - padding)
        end_time = end_sec + padding
        start_frame = int(start_time * sr)
        end_frame = int(end_time * sr)
        sliced = data[start_frame:end_frame]
        
        buffer = io.BytesIO()
        sf.write(buffer, sliced, sr, format='WAV')
        return buffer.getvalue()
    except Exception as e:
        return None

def generate_txt(words_data, include_speakers=False):
    active_words = [w for w in words_data if not w.get("deleted", False)]
    if not include_speakers:
        return " ".join([f"{w.get('prefix_punc', '')}{w['word']}{w.get('punctuation', '')}" for w in active_words])
    txt = ""
    current_speaker = None
    for w in active_words:
        if w["speaker"] != current_speaker:
            spk_name = st.session_state.speaker_names.get(w["speaker"], f"דובר {w['speaker']}")
            txt += f"\n\n[{spk_name}]: "
            current_speaker = w["speaker"]
        txt += f"{w.get('prefix_punc', '')}{w['word']}{w.get('punctuation', '')} "
    return txt.strip()

def generate_srt(words_data):
    active_words = [w for w in words_data if not w.get("deleted", False)]
    srt_content = ""
    chunk_index = 1
    chunk_words = []
    for i, word in enumerate(active_words):
        chunk_words.append(word)
        speaker_changed = i < len(active_words) - 1 and active_words[i+1]['speaker'] != word['speaker']
        if len(chunk_words) >= 10 or speaker_changed or i == len(active_words) - 1:
            start_time = timedelta(seconds=chunk_words[0]['start'])
            end_time = timedelta(seconds=chunk_words[-1]['end'])
            def format_time(td):
                total_sec = int(td.total_seconds())
                hours = total_sec // 3600
                minutes = (total_sec % 3600) // 60
                seconds = total_sec % 60
                millisecs = int((td.total_seconds() - total_sec) * 1000)
                return f"{hours:02d}:{minutes:02d}:{seconds:02d},{millisecs:03d}"
            spk = chunk_words[0]['speaker']
            spk_name = st.session_state.speaker_names.get(spk, f"דובר {spk}")
            speaker_label = f"[{spk_name}] "
            text = " ".join([f"{w.get('prefix_punc', '')}{w['word']}{w.get('punctuation', '')}" for w in chunk_words])
            srt_content += f"{chunk_index}\n{format_time(start_time)} --> {format_time(end_time)}\n{speaker_label}{text}\n\n"
            chunk_index += 1
            chunk_words = []
    return srt_content

def generate_vtt(words_data):
    active_words = [w for w in words_data if not w.get("deleted", False)]
    vtt_content = "WEBVTT\n\n"
    chunk_index = 1
    chunk_words = []
    
    for i, word in enumerate(active_words):
        chunk_words.append(word)
        speaker_changed = i < len(active_words) - 1 and active_words[i+1]['speaker'] != word['speaker']
        if len(chunk_words) >= 10 or speaker_changed or i == len(active_words) - 1:
            start_time = timedelta(seconds=chunk_words[0]['start'])
            end_time = timedelta(seconds=chunk_words[-1]['end'])
            
            def format_time(td):
                total_sec = int(td.total_seconds())
                hours = total_sec // 3600
                minutes = (total_sec % 3600) // 60
                seconds = total_sec % 60
                millisecs = int((td.total_seconds() - total_sec) * 1000)
                if hours > 0:
                    return f"{hours:02d}:{minutes:02d}:{seconds:02d}.{millisecs:03d}"
                return f"{minutes:02d}:{seconds:02d}.{millisecs:03d}"
            
            spk = chunk_words[0]['speaker']
            spk_name = st.session_state.speaker_names.get(spk, f"דובר {spk}")
            speaker_label = f"<v {spk_name}>"
            text = " ".join([f"{w.get('prefix_punc', '')}{w['word']}{w.get('punctuation', '')}" for w in chunk_words])
            
            vtt_content += f"{chunk_index}\n{format_time(start_time)} --> {format_time(end_time)}\n{speaker_label}{text}\n\n"
            chunk_index += 1
            chunk_words = []
    return vtt_content

def generate_docx(words_data):
    if not DOCX_AVAILABLE: return None
    doc = Document()
    doc.add_heading('תמלול שיחה', 0)
    active_words = [w for w in words_data if not w.get("deleted", False)]
    current_speaker = None
    paragraph = None
    for w in active_words:
        if w["speaker"] != current_speaker:
            current_speaker = w["speaker"]
            spk_name = st.session_state.speaker_names.get(current_speaker, f"דובר {current_speaker}")
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"[{spk_name}]: ").bold = True
        paragraph.add_run(f"{w.get('prefix_punc', '')}{w['word']}{w.get('punctuation', '')} ")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# ==========================================
# תפריט צד (Sidebar)
# ==========================================
with st.sidebar:
    st.header("הגדרות")
    st.success("☁️ מחובר לשרתי הענן (API)")
    
    st.divider()
    st.subheader("📁 פרויקטים שמורים")
    saved_projects = [d for d in os.listdir(PROJECTS_DIR) if os.path.isdir(os.path.join(PROJECTS_DIR, d))]
    
    if saved_projects:
        selected_project = st.selectbox(
            "בחר פרויקט מההיסטוריה:", 
            ["-- בחר --"] + saved_projects,
            help="כל תמלול נשמר אוטומטית כפרויקט. בחר פרויקט מהרשימה כדי לטעון את האודיו והטקסט שלו מחדש."
        )
        if selected_project != "-- בחר --":
            proj_path = os.path.join(PROJECTS_DIR, selected_project)
            json_path = os.path.join(proj_path, "data.json")
            try:
                c_time = os.path.getctime(proj_path)
                m_time = os.path.getmtime(json_path) if os.path.exists(json_path) else os.path.getmtime(proj_path)
                c_time_str = datetime.fromtimestamp(c_time).strftime("%d/%m/%Y %H:%M")
                m_time_str = datetime.fromtimestamp(m_time).strftime("%d/%m/%Y %H:%M")
            except Exception:
                c_time_str = "לא ידוע"
                m_time_str = "לא ידוע"
                
            st.caption(f"📅 **תאריך העלאה:** {c_time_str}")
            st.caption(f"✏️ **שינוי אחרון:** {m_time_str}")
            
            if st.button("📂 טען פרויקט", use_container_width=True, type="primary"):
                try:
                    audio_path = os.path.join(proj_path, "audio.wav")
                    if not os.path.exists(json_path):
                        st.error("קובץ data.json חסר בתיקיית הפרויקט.")
                        st.stop()
                    with open(json_path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    
                    st.session_state.original_words_data = data.get("original_words_data", [])
                    st.session_state.words_data = data.get("words_data", [])
                    st.session_state.speaker_names = data.get("speaker_names", {})
                    st.session_state.project_dir = proj_path
                    st.session_state.local_audio_path = audio_path
                    st.session_state.decoded_audio = None 
                    
                    with open(audio_path, "rb") as f:
                        st.session_state.audio_bytes = f.read()
                    
                    st.session_state.active_word_id = ""
                    st.session_state.current_file_name = selected_project
                    st.toast("✅ הפרויקט נטען בהצלחה!")
                    st.rerun()
                except Exception as e:
                    st.error(f"שגיאה בטעינת הפרויקט: {e}")
            
            with st.expander("📝 שינוי שם הפרויקט"):
                new_proj_name = st.text_input("שם חדש:", value=selected_project, key="rename_input", help="תן שם משמעותי לפרויקט כדי שיהיה קל למצוא אותו בעתיד.")
                if st.button("שמור שם", use_container_width=True, help="שומר את שם התיקייה מחדש במחשב."):
                    if new_proj_name and new_proj_name != selected_project:
                        new_proj_path = os.path.join(PROJECTS_DIR, new_proj_name)
                        if os.path.exists(new_proj_path):
                            st.error("⚠️ פרויקט בשם זה כבר קיים!")
                        else:
                            os.rename(proj_path, new_proj_path)
                            if st.session_state.project_dir == proj_path:
                                st.session_state.project_dir = new_proj_path
                                st.session_state.local_audio_path = os.path.join(new_proj_path, "audio.wav")
                                st.session_state.current_file_name = new_proj_name
                            st.success("✅ שם הפרויקט עודכן בהצלחה!")
                            st.rerun()

            with st.expander("🗑️ מחיקת פרויקט"):
                confirm_delete = st.checkbox("כן, אני מבין ומאשר את המחיקה", key="confirm_delete_proj")
                if st.button("מחק פרויקט לצמיתות", type="primary", disabled=not confirm_delete, use_container_width=True, help="מחיקת התיקייה, קובץ השמע ונתוני התמלול. פעולה זו בלתי הפיכה!"):
                    shutil.rmtree(proj_path) 
                    if st.session_state.project_dir == proj_path:
                        st.session_state.words_data = []
                        st.session_state.audio_bytes = None
                        st.session_state.project_dir = None
                    st.success("✅ הפרויקט נמחק בהצלחה!")
                    st.rerun()
    
    st.divider()
    st.subheader("הגדרות תצוגה")
    gap_threshold = st.slider(
        "התרעת השמטה (שניות)", 
        min_value=0.3, max_value=2.0, value=0.6, step=0.1,
        help="מגדיר כמה שניות של שתיקה בין מילים ייחשבו כהפסקה/השמטה. אם יש שתיקה ארוכה מהמוגדר, יופיע סמל ⏳ בעורך הטקסט כדי להתריע על כך."
    )

    st.divider()
    st.subheader("🤖 זיהוי דוברים (AI)")
    
    diarize_option = st.selectbox(
        "בחר מנוע זיהוי דוברים:",
        [
            "Deepgram API (מומלץ לענן - חינמי בהרשמה)",
            "Pyannote בענן (Replicate API - בתשלום סמלי)", 
            "Gemini 1.5 (חינמי לגמרי)"
        ],
        help="לחץ על כל אחת מהאופציות כדי לראות את היתרונות, החסרונות, והעלויות שלה."
    )
    
    expected_speakers = st.number_input("מספר משתתפים צפוי (0 = זיהוי אוטומטי):", min_value=0, max_value=20, value=0, step=1, help="אם אתה יודע כמה אנשים השתתפו בשיחה, הזן את המספר כאן כדי לשפר את הדיוק.")
    
    api_key_diarize = ""
    if "Deepgram" in diarize_option:
        st.info("💡 **Deepgram API:**\n* **יתרון:** מהיר מאוד ומדויק. פתרון אידיאלי לענן.")
        api_key_diarize = st.text_input("מפתח API של Deepgram:", type="password")
        st.caption("[להרשמה והוצאת מפתח מ-Deepgram](https://console.deepgram.com/)")
    elif "Replicate" in diarize_option:
        st.info("💡 **Pyannote 3.1 בענן (Replicate):**\n* **יתרון:** דיוק מילי-שניות מקסימלי.")
        api_key_diarize = st.text_input("מפתח API של Replicate:", type="password")
        st.caption("[להרשמה והוצאת מפתח מ-Replicate](https://replicate.com/account/api-tokens)")
    elif "Gemini" in diarize_option:
        st.info("💡 **Gemini 1.5:**\n* **יתרון:** חינמי לחלוטין דרך Google AI Studio.")
        api_key_diarize = st.text_input("מפתח API של Gemini:", type="password", key="gemini_diarize")
        st.caption("[להוצאת מפתח חינם מ-Google AI Studio](https://aistudio.google.com/app/apikey)")

    if st.session_state.words_data and st.session_state.local_audio_path:
        if st.button("🚀 הפעל זיהוי דוברים עכשיו", use_container_width=True, type="primary"):
            if not api_key_diarize: 
                st.error("חובה להזין מפתח/טוקן מתאים כדי להמשיך.")
            else:
                with st.spinner(f"מנתח קולות באמצעות {diarize_option.split(' ')[0]}..."):
                    try:
                        if "Deepgram" in diarize_option:
                            turns = diarize_with_deepgram(api_key_diarize, st.session_state.local_audio_path, expected_speakers)
                        elif "Replicate" in diarize_option:
                            turns = diarize_with_replicate(api_key_diarize, st.session_state.local_audio_path, expected_speakers)
                        elif "Gemini" in diarize_option:
                            st.warning("זיהוי דוברים דרך Gemini עדיין בפיתוח לגרסה זו.")
                            turns = [] 

                        if turns:
                            for w in st.session_state.words_data:
                                w_start, w_end, max_overlap, assigned_speaker = w["start"], w["end"], 0, "0"
                                for turn in turns:
                                    t_start, t_end, speaker_id = turn["start"], turn["end"], turn["speaker"]
                                    overlap = max(0, min(w_end, t_end) - max(w_start, t_start))
                                    if overlap > max_overlap:
                                        max_overlap, assigned_speaker = overlap, speaker_id
                                    elif t_start > w_end and max_overlap > 0: break
                                w["speaker"] = str(int(assigned_speaker)) if str(assigned_speaker).isdigit() else str(assigned_speaker)
                            
                            save_project()
                            st.success("✅ חלוקת הדוברים בוצעה בהצלחה!")
                            st.rerun()
                            
                    except Exception as e: 
                        st.error(f"שגיאה: {str(e)}")

    st.divider()
    if st.button("איפוס מערכת (Clear Data)", help="מנקה את הזיכרון הזמני של האפליקציה."):
        st.session_state.clear()
        st.rerun()

# ==========================================
# שלב 1: הזנת אודיו ותמלול (מעודכן ל-API וניהול מסלולים)
# ==========================================
st.subheader("1. הזנת אודיו")
col_upload, col_lang = st.columns([6, 4]) 

with col_upload:
    uploaded_file = st.file_uploader(
        "העלה קובץ אודיו (MP3, WAV, M4A, OGG)", 
        type=["mp3", "wav", "m4a", "ogg"],
        help="תומך ברוב קובצי השמע. עובד הכי טוב עם WAV"
    )
    
    st.markdown("<br><br>", unsafe_allow_html=True)
    transcribe_clicked = st.button("🚀 תמלל עכשיו", type="primary", use_container_width=True, help="מתחיל את תהליך זיהוי הדיבור והמרתו לטקסט.")

with col_lang:
    st.write("הגדרות תמלול ומסלול איכות:")
    language_choice = st.radio("שפה", ["he", "en"], format_func=lambda x: "עברית" if x == "he" else "אנגלית", horizontal=True)
    
    engine_tier = st.radio(
        "בחר מסלול תמלול:",
        ["חינמי לחלוטין (Groq)", "איכות פרימיום (כולל ציוני ביטחון)"],
        help="המסלול החינמי מהיר מאוד אך אינו מסמן שגיאות בטקסט. מסלול הפרימיום יאפשר זיהוי מילים בעייתיות בעורך (יושלם בשלב הבא)."
    )
    
    api_key_transcribe = ""
    if engine_tier == "חינמי לחלוטין (Groq)":
        api_key_transcribe = st.secrets.get("GROQ_API_KEY", "")
        if not api_key_transcribe:
            api_key_transcribe = st.text_input("מפתח API של Groq:", type="password", placeholder="gsk_...", help="נדרש מפתח חינמי.")
            st.caption("[להוצאת מפתח חינם בשניות לחץ כאן](https://console.groq.com/keys)")
    else:
        api_key_transcribe = st.secrets.get("PREMIUM_API_KEY", "")
        if not api_key_transcribe:
            api_key_transcribe = st.text_input("מפתח פרימיום (Deepgram/ivrit.ai):", type="password")
    
    st.markdown("💡 **מילון מונחים (Initial Prompt)**")
    initial_prompt = st.text_area(
        "הזן מונחים כאן (מופרדים בפסיקים):", 
        height=80,
        help="המודל לעיתים מתקשה עם שמות פרטיים, סלנג או מונחים מקצועיים. הקלד אותם כאן כדי 'ללמד' אותו."
    )

# טיפול בהעלאת קובץ למערכת
if uploaded_file and (st.session_state.current_file_name != uploaded_file.name):
    st.session_state.words_data = []
    st.session_state.audio_bytes = uploaded_file.getvalue()
    st.session_state.current_file_name = uploaded_file.name
    file_ext = os.path.splitext(uploaded_file.name)[1] or ".wav"
    temp_audio_path = os.path.abspath(os.path.join(TEMP_DIR, f"current_audio{file_ext}"))
    with open(temp_audio_path, "wb") as f: f.write(st.session_state.audio_bytes)
    st.session_state.local_audio_path = temp_audio_path

if transcribe_clicked and st.session_state.audio_bytes:
    if not api_key_transcribe:
        st.error("יש להזין מפתח API מתאים כדי לתמלל בענן.")
    else:
        try:
            with st.spinner("מתמלל בשרתי הענן... (זה לוקח שניות בודדות)"):
                if engine_tier == "חינמי לחלוטין (Groq)":
                    words_list, duration = transcribe_with_groq_api(api_key_transcribe, st.session_state.local_audio_path, language_choice, initial_prompt)
                else:
                    words_list, duration = transcribe_with_premium_api(api_key_transcribe, st.session_state.local_audio_path, language_choice, initial_prompt)
            
            # שמירה והגדרת הפרויקט
            st.session_state.words_data = words_list
            st.session_state.original_words_data = copy.deepcopy(words_list)
            
            project_name = os.path.splitext(st.session_state.current_file_name)[0]
            proj_dir = os.path.join(PROJECTS_DIR, f"{project_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
            os.makedirs(proj_dir, exist_ok=True)
            proj_audio_path = os.path.join(proj_dir, "audio.wav")
            shutil.copy2(st.session_state.local_audio_path, proj_audio_path)
            st.session_state.project_dir = proj_dir
            st.session_state.local_audio_path = proj_audio_path
            save_project()
            st.rerun()
        
        except Exception as e:
            st.error(f"שגיאה בתמלול: {str(e)}")

# ==========================================
# המשך האפליקציה (רק אם יש תמלול)
# ==========================================
if st.session_state.words_data:
    active_words = [w for w in st.session_state.words_data if not w.get("deleted", False)]
    
    # ==========================================
    # שלב 2: ייצוא, סיכום, אוטומציה וסטטיסטיקות
    # ==========================================
    st.divider()
    st.subheader("2. ייצוא, סיכום, פיסוק וסטטיסטיקות")
    
    with st.expander("לחץ כאן לפתיחת אפשרויות הורדה וניתוח AI", expanded=False):
        col_ex1, col_ex2, col_ex3, col_ex4, col_ex5, col_ex6 = st.columns(6)
        final_text = generate_txt(st.session_state.words_data, include_speakers=False)
        final_text_with_speakers = generate_txt(st.session_state.words_data, include_speakers=True)

        with col_ex5: st.download_button("🌐 כתוביות (VTT)", data=generate_vtt(active_words), file_name="subtitles.vtt", mime="text/vtt", use_container_width=True)
        with col_ex1: st.download_button("📝 מתוקן (TXT)", data=final_text_with_speakers, file_name="transcript.txt", mime="text/plain", use_container_width=True)
        with col_ex2: st.download_button("🎬 כתוביות (SRT)", data=generate_srt(active_words), file_name="subtitles.srt", mime="text/plain", use_container_width=True)
        with col_ex3: st.download_button("⚙️ נתונים (JSON)", data=json.dumps(active_words, ensure_ascii=False, indent=2), file_name="data.json", mime="application/json", use_container_width=True)
        with col_ex4:
            if DOCX_AVAILABLE:
                st.download_button("📘 וורד (DOCX)", data=generate_docx(st.session_state.words_data) or "", file_name="transcript.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            else:
                st.button("📘 חסר python-docx", disabled=True, use_container_width=True)

        tab_ai_summary, tab_ai_punc, tab_ai_correction, tab_stats = st.tabs(["🤖 סיכום וניתוח (Gemini)", "✨ פיסוק אוטומטי", "🪄 פרומפט לתיקון ב-AI חיצוני", "📊 סטטיסטיקות"])
        
        # טאב סיכום AI
        with tab_ai_summary:
            st.markdown("הזן מפתח API של Gemini כדי לנתח את הטקסט המתוקן. [הוצאת מפתח בחינם כאן](https://aistudio.google.com/app/apikey)")
            col_gem1, col_gem2 = st.columns([3, 1])
            with col_gem1:
                gemini_key = st.text_input("Gemini API Key:", type="password", key="gemini_key_1")
                prompt_type = st.selectbox("סוג ניתוח:", ["סיכום כללי", "נקודות מפתח (Bullet points)", "ניתוח (אווירה ומסקנות)", "פרומפט מותאם אישית ✍️"])
                custom_prompt_text = st.text_area("הנחיה:", placeholder="לדוגמה: נסח מהשיחה הזו מייל רשמי...") if prompt_type == "פרומפט מותאם אישית ✍️" else ""
            with col_gem2:
                st.write("") 
                st.write("")
                if st.button("🚀 נתח טקסט", type="primary", use_container_width=True):
                    if not gemini_key: st.error("יש להזין מפתח API.")
                    else:
                        with st.spinner("ה-AI מכין תשובה..."):
                            st.session_state.last_ai_response = call_gemini_api(gemini_key, final_text_with_speakers, prompt_type, custom_prompt_text)
            if st.session_state.get("last_ai_response"):
                st.markdown("### ✨ תוצאת הניתוח:")
                with st.container(border=True):
                    st.markdown(st.session_state.last_ai_response)
        
        # טאב פיסוק אוטומטי
        with tab_ai_punc:
            st.markdown("ה-AI יקרא את התמלול, יבין את ההקשר ויוסיף פסיקים ונקודות בדיוק במקומות הנכונים מבלי לפגוע בתזמונים של העורך!")
            col_punc1, col_punc2 = st.columns([3, 1])
            with col_punc1:
                gemini_key_punc = st.text_input("Gemini API Key (לפיסוק):", type="password", key="gemini_key_2")
            with col_punc2:
                st.write("")
                if st.button("✨ פסק אוטומטית", type="primary", use_container_width=True):
                    if not gemini_key_punc:
                        st.error("יש להזין מפתח API.")
                    else:
                        with st.spinner("ה-AI קורא את הטקסט ומחשב סימני פיסוק..."):
                            punct_dict = auto_punctuate_with_gemini(gemini_key_punc, st.session_state.words_data)
                            if punct_dict and isinstance(punct_dict, dict):
                                commit_to_history()
                                count_added = 0
                                for w in st.session_state.words_data:
                                    w_id_str = str(w["id"])
                                    if w_id_str in punct_dict:
                                        w["punctuation"] = punct_dict[w_id_str]
                                        count_added += 1
                                save_project()
                                st.success(f"התווספו סימני פיסוק ל-{count_added} מילים בהצלחה!")
                                st.rerun()
                            else:
                                st.error("ה-AI לא הצליח להחזיר תבנית תקינה. נסה שוב.")
        # טאב פרומפט לתיקון חיצוני (Claude/ChatGPT)
        with tab_ai_correction:
            st.markdown("העתק את הפרומפט הבא והדבק אותו במודל שפה חיצוני (כמו ChatGPT או Claude) יחד עם התמלול המקורי, כדי לתקן שגיאות כתיב פונטיות ביעילות.")
            correction_prompt = """לפניך תמלול גולמי שבוצע על ידי מערכת זיהוי קולי (STT), ולכן יש בו שגיאות כתיב פונטיות. במקביל, אני מצרף קובץ רקע/מילון מושגים הקשור לשיחה. תפקידך הוא לעבור על התמלול ולתקן את שגיאות הכתיב בהתאם למונחים המופיעים בקובץ הרקע.

כללים קריטיים לביצוע:
1. אסור לשנות את תוכן השיחה, לסכם אותה או לשנות את משלב השפה. יש לתקן רק שגיאות זיהוי.
2. חובה לשמור בדיוק על מבנה השורות המקורי.
3. חובה לשמור על תגיות הדוברים בדיוק כפי שהן מופיעות (לדוגמה: [דובר 0]:).

החזר לי את התמלול המתוקן בלבד."""
            
            st.code(correction_prompt, language="text")

        # טאב סטטיסטיקות
        with tab_stats:
            stats = {}
            total_time = active_words[-1]["end"] - active_words[0]["start"] if active_words else 0
            for w in active_words:
                spk, duration = w["speaker"], w["end"] - w["start"]
                if spk not in stats: stats[spk] = {"words": 0, "time_seconds": 0.0}
                stats[spk]["words"] += 1
                stats[spk]["time_seconds"] += duration
                
            cols = st.columns(len(stats))
            for i, (spk, data) in enumerate(stats.items()):
                with cols[i]:
                    spk_name = st.session_state.speaker_names.get(spk, f"דובר {spk}")
                    word_pct = (data["words"] / len(active_words)) * 100 if active_words else 0
                    st.markdown(f"### {spk_name}")
                    st.metric("סה״כ מילים", f'{data["words"]} ({word_pct:.1f}%)')
                    st.metric("זמן דיבור נטו", f'{data["time_seconds"]/60:.1f} דק׳')

    # ==========================================
    # שלב 3: כלי חיפוש ומיזוג דוברים
    # ==========================================
    st.divider()
    st.subheader("3. חיפוש, החלפה וניהול דוברים")
    
    col_search, col_replace, col_speakers = st.columns([1, 1, 1.5])
    
    with col_search:
        search_term = st.text_input("🔍 הדגשת מילה בעורך:", "", help="הקלד מילה כדי לסמן אותה בירוק לאורך כל התמלול בעורך למטה.")
        
    with col_replace:
        with st.form("replace_form"):
            rep_search = st.text_input("החלפה גורפת (חפש):", help="המילה שתרצה למצוא ולתקן בכל התמלול.")
            rep_target = st.text_input("החלף ב:", help="המילה התקינה שתחליף את המילה השגויה.")
            if st.form_submit_button("🔄 החלף בכל הטקסט", use_container_width=True, help="מחליף בבת אחת את כל המופעים של המילה ושומר להיסטוריית הביטולים (Undo)."):
                if rep_search:
                    commit_to_history()
                    count = 0
                    for w in st.session_state.words_data:
                        if not w.get("deleted") and (w["word"] == rep_search or w["clean_word"] == rep_search):
                            w["word"] = rep_target
                            w["confidence"] = 1.0
                            count += 1
                    if count > 0:
                        save_project()
                        st.success(f"הוחלפו {count} מופעים!")
                        st.rerun()
                    else:
                        st.warning("המילה לא נמצאה.")
        
    with col_speakers:
        with st.expander("👥 מיזוג ושינוי שמות דוברים (Bulk)", expanded=False):
            st.caption("כדי למזג, פשוט תן לשני דוברים את אותו השם.")
            unique_speakers = sorted(list(set([w["speaker"] for w in active_words])))
            with st.form("bulk_speaker_management"):
                new_names = {}
                for spk in unique_speakers:
                    current_name = st.session_state.speaker_names.get(spk, spk)
                    new_names[spk] = st.text_input(f"דובר מקורי [{spk}]:", value=current_name)
                    
                if st.form_submit_button("💾 שמור ומזג דוברים", type="primary", use_container_width=True):
                    commit_to_history()
                    for w in st.session_state.words_data:
                        if w["speaker"] in new_names:
                            new_spk_name = new_names[w["speaker"]].strip()
                            if new_spk_name: w["speaker"] = new_spk_name
                    st.session_state.speaker_names = {name.strip(): name.strip() for name in new_names.values() if name.strip()}
                    save_project()
                    st.rerun()

    # ==========================================
    # שלב 4: עריכת תמלול (העורך המרכזי)
    # ==========================================
    st.divider()
    
    col_title, col_speed, col_undo, col_redo = st.columns([5, 1.5, 1, 1])
    with col_title:
        st.subheader("4. עריכת תמלול (הקלדה חופשית)")
        st.caption("✨ **קיצורים:** ⏸️ **Ctrl+Space** (ניגון) | ↩️ **Ctrl+Z** (בטל) | ↪️ **Ctrl+Y** (שחזר) | 🔄 **לחיצה כפולה** (שינוי דובר) | ⏱️ **לחיצה על הזמן** (קפיצה)")
    
    with col_speed:
        playback_speed = st.selectbox("⚡ מהירות נגן:", [0.75, 1.0, 1.25, 1.5, 2.0], index=1)

    with col_undo:
        st.write("") 
        if st.button("↩️ בטל", disabled=(st.session_state.history_index <= 0), use_container_width=True):
            st.session_state.history_index -= 1
            st.session_state.words_data = copy.deepcopy(st.session_state.history[st.session_state.history_index])
            save_project()
            st.rerun()
            
    with col_redo:
        st.write("") 
        if st.button("↪️ שחזר", disabled=(st.session_state.history_index >= len(st.session_state.history) - 1), use_container_width=True):
            st.session_state.history_index += 1
            st.session_state.words_data = copy.deepcopy(st.session_state.history[st.session_state.history_index])
            save_project()
            st.rerun()

    # הוספת הנגן המקורי של Streamlit 
    if st.session_state.get("local_audio_path") and os.path.exists(st.session_state.local_audio_path):
        st.audio(st.session_state.local_audio_path)

    # קריאה לרכיב ה-React
    component_value = custom_transcription_editor(
        words_data=st.session_state.words_data,
        speaker_names=st.session_state.speaker_names,
        gap_threshold=gap_threshold,
        search_query=search_term,
        playback_rate=playback_speed,
        key="main_editor"
    )
    
    # עריכת מילה ספציפית (לחיצה כפולה)
    if st.session_state.active_word_id:
        selected_id = int(st.session_state.active_word_id)
        selected_word = next((w for w in st.session_state.words_data if w["id"] == selected_id), None)
        
        if selected_word:
            with st.container(border=True):
                st.markdown(f"🎯 **עריכת דובר למילה:** `{selected_word['word']}`")
                
                col_audio, col_speaker = st.columns([1, 1])
                with col_audio:
                    sliced = slice_audio(selected_word["start"], selected_word["end"], padding=1.0)
                    if sliced:
                        b64 = base64.b64encode(sliced).decode()
                        audio_html = f"""
                        <div style="display:flex; justify-content:center; align-items:center; height:100%;">
                            <audio controls autoplay style="width: 100%; height: 45px; border-radius:8px;">
                              <source src="data:audio/wav;base64,{b64}" type="audio/wav">
                            </audio>
                        </div>
                        """
                        components.html(audio_html, height=65)
                
                with col_speaker:
                    with st.form(key=f"speaker_form_{selected_id}"):
                        c1, c2 = st.columns([2, 1])
                        with c1:
                            new_spk = st.text_input("הגדר דובר חדש למילה זו:", value=selected_word["speaker"])
                        with c2:
                            st.write("") 
                            st.write("")
                            if st.form_submit_button("שמור דובר", type="primary", use_container_width=True):
                                idx = st.session_state.words_data.index(selected_word)
                                st.session_state.words_data[idx]["speaker"] = new_spk
                                save_project()
                                st.session_state.active_word_id = "" 
                                st.rerun()

    # ניהול נתונים שחוזרים מ-React
    if component_value and isinstance(component_value, dict):
        action = component_value.get("action")
        ts = component_value.get("ts") 
        
        if ts and ts != st.session_state.get("last_processed_ts"):
            st.session_state.last_processed_ts = ts 
            
            if action == "update":
                new_data = component_value.get("data")
                if new_data:
                    commit_to_history()
                    st.session_state.words_data = new_data
                    save_project()
                    st.rerun()
                    
            elif action == "select":
                selected_id = str(component_value.get("word_id"))
                st.session_state.active_word_id = selected_id
                st.rerun()
            elif action == "undo":
                if st.session_state.history_index > 0:
                    st.session_state.history_index -= 1
                    st.session_state.words_data = copy.deepcopy(st.session_state.history[st.session_state.history_index])
                    save_project()
                    st.rerun()
            elif action == "redo":
                if st.session_state.history_index < len(st.session_state.history) - 1:
                    st.session_state.history_index += 1
                    st.session_state.words_data = copy.deepcopy(st.session_state.history[st.session_state.history_index])
                    save_project()
                    st.rerun()